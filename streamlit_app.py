import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt

# Function to apply styles to the text
def set_paragraph_style(paragraph, font_size):
    run = paragraph.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)

# Mock humanizer function (for demonstration)
def humanize_content(text):
    # In a real implementation, this function would process the text through an AI humanizer
    # For now, let's just return the text as is for demonstration purposes.
    return text

# Function to generate and save the essay
def generate_and_save_essay(topic, api_key):
    # Configure the Google Generative AI API
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-1.5-flash")

    # Generate content based on the user-defined topic
    response = model.generate_content(f"Write a detailed essay on the topic: {topic}")

    if response.text:
        # Humanize the content
        humanized_content = humanize_content(response.text)

        # Create a new Document
        doc = Document()

        # Title and Description
        doc.add_heading("Essay Generator with Google Generative AI", level=1)
        set_paragraph_style(doc.paragraphs[-1], 15)  # Header size 15
        doc.add_paragraph("This Streamlit app allows users to generate detailed essays using Google Generative AI (Gemini). Users can input their desired topic and receive a comprehensive essay generated by the AI.", style='Normal')
        set_paragraph_style(doc.paragraphs[-1], 11)  # Normal text size 11

        # Adding humanized essay content
        doc.add_heading("Generated Essay:", level=2)
        set_paragraph_style(doc.paragraphs[-1], 14)  # Sub-header size 14
        doc.add_paragraph(humanized_content, style='Normal')
        set_paragraph_style(doc.paragraphs[-1], 11)  # Normal text size 11

        # Save the document
        file_path = "/mnt/data/Essay_Generator_Output.docx"
        doc.save(file_path)
        return file_path
    else:
        return None

# Streamlit App Title and Description
st.title("Essay Generator with Google Generative AI")
st.write("Enter a topic, and this app will generate a detailed essay for you.")

# Input field for API key and topic
api_key = "AIzaSyBzP_urPbe1zBnZwgjhSlVl-MWtUQMEqQA"
topic = st.text_input("Enter the topic for the essay")

# Button to trigger essay generation
if st.button("Generate Essay"):
    if not api_key or not topic:
        st.error("Please make sure to enter all fields: API Key and Topic.")
    else:
        output_file_path = generate_and_save_essay(topic, api_key)

        if output_file_path:
            st.success("Essay generated successfully!")
            st.download_button(label="Download the essay as DOCX", data=open(output_file_path, "rb"), file_name="Essay_Generator_Output.docx")
        else:
            st.error("Failed to generate the essay. Please try again.")
